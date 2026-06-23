import pandas as pd
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def fix_capitalization(name):
    if pd.isna(name) or name == '':
        return name
    name = str(name).strip()
    if name.isupper() or name.islower():
        return name.title()
    return name

def parse_name_with_suffix(full_name, default_last_name=''):
    """Parse names, handling Jr/Sr/II/III/IV suffixes properly."""
    if pd.isna(full_name) or full_name == '':
        return None, None, False

    full_name = str(full_name).strip()
    full_name = full_name.replace('*', '')
    full_name = re.sub(r'\([^)]*\)', '', full_name).strip()

    # Remove relationship words
    full_name = re.split(r'\s*[-–—]\s*(?:son|daughter|grandson|granddaughter|parent|mother|father|sister|brother|friend|guest)\s*$', full_name, flags=re.IGNORECASE)[0]
    full_name = full_name.strip()

    parts = full_name.split()
    if len(parts) == 0:
        return None, None, False
    elif len(parts) == 1:
        return fix_capitalization(parts[0]), fix_capitalization(default_last_name) if default_last_name else default_last_name, True

    # Check for suffix
    suffix_pattern = r'^(Jr\.?|Sr\.?|II|III|IV|V)$'
    has_suffix = len(parts) >= 2 and re.match(suffix_pattern, parts[-1], re.IGNORECASE)

    if has_suffix and len(parts) >= 3:
        first_name = ' '.join([fix_capitalization(p) for p in parts[:-2]])
        last_name = fix_capitalization(parts[-2]) + ' ' + parts[-1]
        return first_name, last_name, False
    elif len(parts) == 2:
        # Could be "FirstName LastName" or "LastName Jr" - assume first case
        first_name = fix_capitalization(parts[0])
        last_name = fix_capitalization(parts[1])
        return first_name, last_name, False
    else:
        first_name = ' '.join([fix_capitalization(p) for p in parts[:-1]])
        last_name = fix_capitalization(parts[-1])
        return first_name, last_name, False

def parse_guests_string(guests_str):
    if pd.isna(guests_str) or guests_str == '':
        return []
    guests_str = str(guests_str)
    guests_str = re.sub(r'\s+and\s+', ', ', guests_str)
    guests_str = re.sub(r'\n+', ', ', guests_str)
    names = [name.strip() for name in guests_str.split(',')]
    return [name for name in names if name and len(name) > 1]

def text_to_number(text):
    if pd.isna(text):
        return 0
    text = str(text).lower().strip()
    word_to_num = {'one': 1, 'two': 2, 'three': 3, 'four': 4, 'five': 5,
                   'six': 6, 'seven': 7, 'eight': 8, 'nine': 9, 'ten': 10}
    for word, num in word_to_num.items():
        if word in text:
            return num
    match = re.search(r'\d+', text)
    if match:
        return int(match.group())
    return 0

# Load base Culbreth data
print("Loading base Culbreth data...")
culbreth_records = []
existing = pd.read_excel('C:/Users/Ben/Downloads/Graduation_Parking_Reconciled_2026_v2.xlsx')
for idx, row in existing.iterrows():
    last_name = str(row['Last Name']).replace('*', '').strip()
    last_name = re.sub(r'\([^)]*\)', '', last_name).strip()
    first_name = str(row['First Name']).replace('*', '').strip()
    first_name = re.sub(r'\([^)]*\)', '', first_name).strip()

    if last_name and last_name != 'nan':
        culbreth_records.append({
            'Last Name': last_name,
            'First Name': first_name,
            'Affiliation': row['Affiliation'],
            'Ceremony Day': row['Ceremony Day'],
            'Number of Passes': row['Number of Passes'],
            'Notes': ''
        })
print(f"  Loaded {len(culbreth_records)} records")

# Add BOV
print("\nAdding BOV...")
bov_df = pd.read_excel('C:/Users/Ben/Downloads/BOV_Participation_Matrix_2026.xlsx',
                       sheet_name='Participation Matrix', header=0)
bov_count = 0
for idx, row in bov_df.iterrows():
    if idx == 0:
        continue
    last_name = row.get('Unnamed: 1', '')
    status = row.get('Unnamed: 2', '')
    guest_names_str = row.get('Unnamed: 10', '')

    if pd.isna(last_name) or last_name == '' or status == 'Not participating':
        continue

    guest_names = parse_guests_string(guest_names_str)
    for guest_name in guest_names:
        if any(x in guest_name.lower() for x in ['saturday', 'sunday', 'valedictory']):
            continue
        first, last, used_default = parse_name_with_suffix(guest_name, default_last_name=last_name)
        if first and last:
            culbreth_records.append({
                'Last Name': last,
                'First Name': first,
                'Affiliation': 'BOV',
                'Ceremony Day': 'Both',
                'Number of Passes': 1,
                'Notes': ''
            })
            bov_count += 1
print(f"  Added {bov_count} BOV guests")

# Add Credentials
print("\nAdding Credentials Master List...")
creds_df = pd.read_excel('C:/Users/Ben/Downloads/Credentials Master List -  2026 Final Exercises as of 4.30.2026.xlsx')
creds_count = 0
for idx, row in creds_df.iterrows():
    first_name = fix_capitalization(row.get('First Name', ''))
    last_name = fix_capitalization(row.get('Last Name', ''))
    department = row.get('Department', '')
    title = row.get('Title', '')

    last_name = last_name.replace('*', '').strip() if pd.notna(last_name) else ''
    first_name = first_name.replace('*', '').strip() if pd.notna(first_name) else ''

    culbreth_passes = row.get('Culbreth Garage / \nJPJ Garage', 0)
    try:
        culbreth_passes = float(culbreth_passes) if pd.notna(culbreth_passes) else 0
    except:
        culbreth_passes = 0

    if culbreth_passes <= 0 or not last_name:
        continue

    affiliation = str(department).strip() if pd.notna(department) and department != '' else (
        str(title).strip() if pd.notna(title) and title != '' else 'Special Guest')

    culbreth_records.append({
        'Last Name': last_name,
        'First Name': first_name,
        'Affiliation': affiliation,
        'Ceremony Day': 'Both',
        'Number of Passes': int(culbreth_passes),
        'Notes': ''
    })
    creds_count += 1
print(f"  Added {creds_count} from Credentials")

# Process Platform & Procession for Carr's Hill
print("\nProcessing Platform & Procession Party...")
carrs_hill_records = []
platform_df = pd.read_excel('C:/Users/Ben/Downloads/Finals Weekend 2026 - Platform & Procession Party_4.15.2026.xlsx')

for idx, row in platform_df.iterrows():
    first_name = fix_capitalization(row.get('First Name:', ''))
    last_name = fix_capitalization(row.get('Last Name:', ''))
    title = row.get('Title and School or Department:', '')

    if pd.isna(last_name) or last_name == '':
        continue

    affiliation = str(title).strip() if pd.notna(title) and title != '' else 'Platform Party'

    # Check "Both?" column
    both = row.get('Both?', '')
    sat_passes_text = row.get('Seating passes \nSaturday, May 16th', '')
    sun_passes_text = row.get('Seating passes \nSunday, May 17th', '')

    if both == 'X':
        # Both days - combine pass counts
        sat_passes = text_to_number(sat_passes_text) if pd.notna(sat_passes_text) else 0
        sun_passes = text_to_number(sun_passes_text) if pd.notna(sun_passes_text) else 0
        total_passes = max(sat_passes, sun_passes, 1)  # At least 1 pass

        carrs_hill_records.append({
            'Last Name': last_name,
            'First Name': first_name,
            'Affiliation': affiliation,
            'Ceremony Day': 'Both',
            'Number of Passes': total_passes,
            'Notes': ''
        })
    else:
        # Individual days
        sat_passes = text_to_number(sat_passes_text)
        sun_passes = text_to_number(sun_passes_text)

        if sat_passes > 0:
            carrs_hill_records.append({
                'Last Name': last_name,
                'First Name': first_name,
                'Affiliation': affiliation,
                'Ceremony Day': 'Saturday',
                'Number of Passes': sat_passes,
                'Notes': ''
            })

        if sun_passes > 0:
            carrs_hill_records.append({
                'Last Name': last_name,
                'First Name': first_name,
                'Affiliation': affiliation,
                'Ceremony Day': 'Sunday',
                'Number of Passes': sun_passes,
                'Notes': ''
            })

print(f"  Added {len(carrs_hill_records)} Platform Party records")

# Deduplicate
def affiliation_priority(affiliation):
    return 0 if affiliation == 'Special Guest' else 1

df_culbreth = pd.DataFrame(culbreth_records)
df_culbreth['affiliation_priority'] = df_culbreth['Affiliation'].apply(affiliation_priority)
df_culbreth = df_culbreth.sort_values(
    by=['Last Name', 'First Name', 'affiliation_priority', 'Number of Passes'],
    ascending=[True, True, False, False]
)
df_culbreth = df_culbreth.drop_duplicates(subset=['Last Name', 'First Name'], keep='first')
df_culbreth = df_culbreth.drop(columns=['affiliation_priority'])
df_culbreth = df_culbreth.sort_values(by=['Last Name', 'First Name'])

# Remove standalone Jr/Sr entries (incomplete names)
df_culbreth = df_culbreth[~df_culbreth['Last Name'].isin(['Jr', 'Jr.', 'Sr', 'Sr.'])]

df_carrs = pd.DataFrame(carrs_hill_records)
df_carrs['affiliation_priority'] = df_carrs['Affiliation'].apply(affiliation_priority)
df_carrs = df_carrs.sort_values(
    by=['Last Name', 'First Name', 'affiliation_priority', 'Number of Passes'],
    ascending=[True, True, False, False]
)
df_carrs = df_carrs.drop_duplicates(subset=['Last Name', 'First Name'], keep='first')
df_carrs = df_carrs.drop(columns=['affiliation_priority'])
df_carrs = df_carrs.sort_values(by=['Last Name', 'First Name'])

# Save Excel
with pd.ExcelWriter('C:/Users/Ben/Downloads/Graduation_Parking_Reconciled_2026_Final.xlsx', engine='openpyxl') as writer:
    df_culbreth.to_excel(writer, sheet_name='Culbreth Garage', index=False)
    df_carrs.to_excel(writer, sheet_name='Carrs Hill Madison Hall', index=False)

print(f"\nFinal counts:")
print(f"  Culbreth Garage: {len(df_culbreth)}")
print(f"  Carrs Hill / Madison Hall: {len(df_carrs)}")

# Create Word docs
def create_word_doc(df, location_name, output_file):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    title = doc.add_paragraph()
    title_run = title.add_run(f'{location_name} Parking List - Final Exercises 2026')
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date = doc.add_paragraph()
    date_run = date.add_run('Sat. May 17 & Sun. May 18, 2026')
    date_run.font.size = Pt(12)
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'

    widths = [Inches(0.8), Inches(1.0), Inches(1.3), Inches(1.3), Inches(1.5), Inches(1.0)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    header_cells = table.rows[0].cells
    header_cells[0].text = 'Day'
    header_cells[1].text = 'Location'
    header_cells[2].text = 'First Name'
    header_cells[3].text = 'Last Name'
    header_cells[4].text = 'Affiliation'
    header_cells[5].text = 'Spaces/Notes'

    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for idx, row in df.iterrows():
        row_cells = table.add_row().cells

        for idx_cell, width in enumerate(widths):
            row_cells[idx_cell].width = width

        row_cells[0].text = str(row['Ceremony Day'])
        row_cells[1].text = location_name
        row_cells[2].text = str(row['First Name'])
        row_cells[3].text = str(row['Last Name'])
        row_cells[4].text = str(row['Affiliation'])

        spaces = str(int(row['Number of Passes'])) if row['Number of Passes'] > 0 else ''
        row_cells[5].text = spaces

        for cell in row_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    doc.save(output_file)

create_word_doc(df_culbreth, 'Culbreth', 'C:/Users/Ben/Downloads/Culbreth_Parking_List_2026.docx')
create_word_doc(df_carrs, "Carr's Hill / Madison Hall", 'C:/Users/Ben/Downloads/Carrs_Hill_Parking_List_2026.docx')

print("\nAll files created successfully!")
